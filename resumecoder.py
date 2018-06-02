import os
import yaml

from hax import add_hyperlink
from docx import Document


def islist(thing):
    return isinstance(thing, list)


def isdict(thing):
    return isinstance(thing, dict)


def isstr(thing):
    return isinstance(thing, str)


def yml_to_dict(path: str) -> dict:
    with open(path, 'r') as f:
        return yaml.safe_load(f)


class ResumeCoder:
    document = Document()
    path = ""
    data = {}

    def __init__(self, path):
        self.path = path

    def write(self, path="resume.docx"):
        """Save document to file."""
        self.document.save(os.path.join(self.path, path))

    def contact_info(self, path: str, sep=" | "):

        def name(name, doc):
            doc.add_paragraph(name)

        def email(email, doc):

            p = doc.add_paragraph('')

            if isstr(email):  # just one email
                p.add_run(email)

            elif islist(email):  # list of emails
                for i in range(len(email)):
                    p.add_run(email[i])

                    if i < len(email) - 1:  # prevent sep at end of list
                        p.add_run(sep)

            elif isdict(email):  # emails labeled as 'home', 'work', etc
                i = 0

                for key, value in email.items():
                    h = add_hyperlink(p, text=value, url=f'mailto:{value}')
                    p.add_run(' ' + '(' + key + ')')

                    if i < (len(email.items()) - 1):  # prevent sep at end
                        p.add_run(sep)
                    i += 1

        def address(address, doc):
            p = doc.add_paragraph('')

            for key, value in address.items():
                print(key, value)

                p.add_run(value['street'] + ", ")
                p.add_run(value['city'] + ", ")
                p.add_run(value['state'] + ", ")
                p.add_run(str(value['zip']))

                p.add_run(f' ({key})')

        def phone(phone, doc):
            p = doc.add_paragraph('')

            if isdict(phone):
                for key, value in phone.items():
                    p.add_run(value)
                    p.add_run(f" ({key})")

            elif isstr(phone):
                p.add_run(phone)

        path = os.path.join(self.path, path)
        data = yml_to_dict(path)

        name(data['name'], self.document)

        if 'email' in data:
            email(data['email'], self.document)

        if 'address' in data:
            address(data['address'], self.document)

        if 'phone' in data:
            phone(data['phone'], self.document)


if __name__ == '__main__':
    rc = ResumeCoder('./data/Henry/')
    rc.contact_info('contact info.yml')
    rc.write()
